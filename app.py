import streamlit as st
from groq import Groq
from pypdf import PdfReader
from duckduckgo_search import DDGS
from docx import Document
from io import BytesIO

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="AEQUITAS AI | Legal Intelligence", 
    page_icon="⚖️", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- DESIGN CUSTOMIZADO (CSS) ---
st.markdown("""
    <style>
    /* Estética Geral */
    .main { background-color: #0e1117; color: #ffffff; }
    
    /* Estilo dos Botões */
    .stButton>button {
        width: 100%;
        border-radius: 8px;
        height: 3em;
        background-color: #c9a44c; /* Dourado Jurídico */
        color: white;
        border: none;
        font-weight: bold;
        transition: 0.3s;
    }
    .stButton>button:hover {
        background-color: #a68539;
        border: 1px solid #ffffff;
    }
    
    /* Caixas de Mensagem */
    .stChatMessage {
        background-color: #1a1c24;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 15px;
        border: 1px solid #2d2f3b;
    }
    
    /* Sidebar */
    [data-testid="stSidebar"] {
        background-color: #161b22;
        border-right: 1px solid #30363d;
    }
    
    /* Títulos */
    h1, h2, h3 { 
        color: #c9a44c; 
        font-family: 'Georgia', serif; 
    }
    
    /* Esconder o menu padrão do Streamlit para parecer App profissional */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    </style>
    """, unsafe_allow_html=True)

# --- INICIALIZAÇÃO DE SESSÃO ---
if "chats" not in st.session_state:
    st.session_state.chats = {"Nova Consulta": []}
if "current_chat" not in st.session_state:
    st.session_state.current_chat = "Nova Consulta"
if "last_response" not in st.session_state:
    st.session_state.last_response = ""

# Função para gerar Word
def criar_docx(texto):
    doc = Document()
    doc.add_heading('PARECER TÉCNICO - AEQUITAS AI', 0)
    doc.add_paragraph(texto)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- BARRA LATERAL (BRANDING & CONTROLOS) ---
with st.sidebar:
    st.markdown("<h1 style='text-align: center;'>⚖️</h1>", unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: center; margin-top: -20px;'>AEQUITAS AI</h2>", unsafe_allow_html=True)
    st.caption("<p style='text-align: center;'>Equidade e Precisão no Direito Português</p>", unsafe_allow_html=True)
    st.markdown("---")
    
    if st.button("➕ NOVO PROCESSO"):
        nome_novo = f"Processo {len(st.session_state.chats) + 1}"
        st.session_state.chats[nome_novo] = []
        st.session_state.current_chat = nome_novo
        st.rerun()

    st.subheader("📁 Histórico de Casos")
    escolha = st.radio("Selecione a conversa:", list(st.session_state.chats.keys()))
    st.session_state.current_chat = escolha
    
    st.markdown("---")
    st.subheader("⚙️ Painel de Controlo")
    api_key = st.text_input("Groq API Key:", type="password", help="Insira a sua chave para ativar o cérebro da IA.")
    perfil = st.selectbox("Especialidade:", 
                         ["Consultor Geral", "Analisador de Acórdãos", "Especialista em Direito Civil", 
                          "Especialista em Direito Penal", "Revisor de Estilo Jurídico"])
    
    usar_web = st.checkbox("Pesquisa Web Ativa (DRE/PGDL)")
    uploaded_file = st.file_uploader("Upload de Documento Base (PDF)", type="pdf")

# --- ÁREA PRINCIPAL ---
st.title(f"📄 {st.session_state.current_chat}")
st.info(f"Foco atual: {perfil}")

# Processamento de PDF
contexto_pdf = ""
if uploaded_file:
    reader = PdfReader(uploaded_file)
    contexto_pdf = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    st.toast("PDF integrado com sucesso!", icon="📚")

# Mostrar histórico de mensagens
for message in st.session_state.chats[st.session_state.current_chat]:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Entrada de Chat
if prompt := st.chat_input("Como posso ajudar na sua estratégia jurídica hoje?"):
    if not api_key:
        st.error("⚠️ Por favor, insira a sua API Key na barra lateral para continuar.")
    else:
        # Adicionar pergunta do utilizador
        st.session_state.chats[st.session_state.current_chat].append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        client = Groq(api_key=api_key)
        
        # Pesquisa Web
        contexto_web = ""
        if usar_web:
            with st.spinner("A consultar fontes oficiais portuguesas..."):
                try:
                    with DDGS() as ddgs:
                        busca = f"{prompt} legislação portugal site:dre.pt OR site:pgdlisboa.pt"
                        res = ddgs.text(busca, max_results=3)
                        contexto_web = "\n".join([r['body'] for r in res])
                except: pass

        # Prompt de Sistema Refinado
        sistema_prompt = f"""
        És o AEQUITAS AI, um assistente jurídico de elite em Portugal.
        PERFIL: {perfil}.
        OBJETIVO: Ser técnico, preciso e honesto. 
        CONTEXTO PDF: {contexto_pdf[:7000]}
        CONTEXTO WEB: {contexto_web}
        REGRA: Nunca inventes artigos. Se não souberes, sugere a consulta direta no Diário da República.
        """

        with st.chat_message("assistant"):
            # Histórico limitado para performance
            historico = st.session_state.chats[st.session_state.current_chat][-6:]
            mensagens_envio = [{"role": "system", "content": sistema_prompt}] + historico
            
            try:
                completion = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=mensagens_envio,
                    temperature=0.1
                )
                response = completion.choices[0].message.content
                st.markdown(response)
                
                # Guardar resposta
                st.session_state.chats[st.session_state.current_chat].append({"role": "assistant", "content": response})
                
                # Botão de Download direto na resposta
                st.download_button(
                    label="📥 Exportar Parecer para Word",
                    data=criar_docx(response),
                    file_name=f"Parecer_{st.session_state.current_chat}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Erro de conexão: {e}. Tente reduzir o tamanho do texto ou aguardar alguns minutos.")

# Botão para apagar chat na barra lateral
if st.sidebar.button("🗑️ APAGAR ESTE PROCESSO"):
    if len(st.session_state.chats) > 1:
        del st.session_state.chats[st.session_state.current_chat]
        st.session_state.current_chat = list(st.session_state.chats.keys())[0]
        st.rerun()
    else:
        st.session_state.chats = {"Nova Consulta": []}
        st.rerun()
